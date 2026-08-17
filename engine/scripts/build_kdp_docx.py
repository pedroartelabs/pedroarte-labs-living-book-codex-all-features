"""Constrói deterministicamente o DOCX de interior para a Amazon KDP.

Este script generaliza `scripts/build_kdp_docx.py` do runtime
`a_morte_ainda_nao_nasceu`, cujo resultado passou em 71 PASS / 0 FAIL na
auditoria estrutural e teve as 263 páginas renderizadas inspecionadas uma a
uma. A lógica editorial é a mesma; o que muda é que nada é hardcoded: título,
autor, capítulos, tipografia, geometria e posicionamento de imagem vêm de
arquivos de configuração.

Por que isto vive no motor, e não no runtime de cada livro
-----------------------------------------------------------
Na execução real do Codex, esta capacidade foi reconstruída do zero dentro do
runtime — mais de 1.100 linhas de ferramenta para um problema que é idêntico
em qualquer livro composto por este motor (mesmas seções, mesmas margens
espelhadas, mesmo sumário com PAGEREF, mesmo placement OPEN/HINGE). Reescrever
isso a cada obra contraria o princípio central do repositório: o motor cuida da
fábrica, o pacote define a criatura. Além disso, montar um DOCX é trabalho
inteiramente determinístico — não há julgamento criativo em posicionar uma
quebra de seção. Fazer isso com um agente de linguagem gasta tokens e tempo
para produzir algo que um script produz igual, sempre, de graça.

Melhoria embutida em relação ao original
----------------------------------------
O build original usava um estilo customizado `Chapter Title` e precisava de um
script de pós-processamento (`apply_heading1_to_chapters.py`) para converter os
títulos ao `Heading 1` interno, de modo que o Word conseguisse gerar sumário
automático. Aqui os títulos já nascem em `Heading 1`, com a mesma aparência
(Georgia Bold 18 pt, centralizado, keep-with-next, 24 pt depois) e nível de
estrutura 0. Um script inteiro deixa de ser necessário.

Entradas (relativas à raiz do runtime)
--------------------------------------
- `book/BOOK_SPEC.yaml`        título, autor, idioma, capítulos e seus títulos
- `layout/KDP_LAYOUT.yaml`     opcional; sobrescreve os defaults do motor
- `layout/IMAGE_PLACEMENT.yaml` opcional; placement/âncora/alt text por capítulo
- `manuscript/final/...md`     prosa congelada (caminho vem da config)
- `images/approved/*.jpg`      imagens aprovadas

Saída: `outputs/KDP_DRAFT.docx` (caminho configurável).

Uso:
    python engine/scripts/build_kdp_docx.py --runtime runtime/<slug>
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _layout_config import (  # noqa: E402
    load_book_metadata as load_book,
    load_layout_config as load_config,
    load_yaml,
)

# Placements aceitos. OPEN ocupa o verso anterior à abertura do capítulo;
# HINGE e AFTER só entram depois do parágrafo-gatilho declarado.
PLACEMENT_BEFORE = {"OPEN"}
PLACEMENT_AFTER = {"HINGE", "AFTER"}


def load_placement(runtime: Path, chapter_count: int) -> dict[int, dict]:
    """Placement por capítulo. Sem o arquivo, todo capítulo abre com imagem
    (OPEN) — o comportamento mais conservador, porque OPEN nunca antecipa um
    gatilho narrativo que ainda não aconteceu."""
    path = runtime / "layout" / "IMAGE_PLACEMENT.yaml"
    if not path.is_file():
        return {c: {"placement": "OPEN"} for c in range(1, chapter_count + 1)}
    raw = load_yaml(path).get("chapters", {}) or {}
    out: dict[int, dict] = {}
    for key, value in raw.items():
        entry = dict(value or {})
        entry.setdefault("placement", "OPEN")
        out[int(key)] = entry
    return out


# --------------------------------------------------------------------------
# Helpers OOXML
# --------------------------------------------------------------------------

def set_run_font(run, family, language, size=None, bold=None, italic=None):
    run.font.name = family
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:ascii"), family)
    rpr.rFonts.set(qn("w:hAnsi"), family)
    rpr.rFonts.set(qn("w:cs"), family)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), language)


def ensure_style(doc, name, family, *, size, bold=False, alignment=None, before=0,
                 after=0, line=None, first=0.0, left=0.0, right=0.0,
                 keep_next=False, widow=True, outline=None):
    styles = doc.styles
    if name in styles and styles[name].type != WD_STYLE_TYPE.PARAGRAPH:
        styles.element.remove(styles[name]._element)
    style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = family
    rpr = style._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:ascii"), family)
    rpr.rFonts.set(qn("w:hAnsi"), family)
    style.font.size = Pt(size)
    style.font.bold = bold
    fmt = style.paragraph_format
    if alignment is not None:
        fmt.alignment = alignment
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.first_line_indent = Inches(first)
    fmt.left_indent = Inches(left)
    fmt.right_indent = Inches(right)
    fmt.keep_with_next = keep_next
    fmt.widow_control = widow
    if line:
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = Pt(line)
    if outline is not None:
        # Nível de estrutura: é o que permite ao Word montar sumário automático.
        ppr = style._element.get_or_add_pPr()
        lvl = ppr.find(qn("w:outlineLvl"))
        if lvl is None:
            lvl = OxmlElement("w:outlineLvl")
            ppr.append(lvl)
        lvl.set(qn("w:val"), str(outline))
    return style


def setup_styles(doc, typo):
    family = typo["family"]
    # Word traz "Book Title" e "Page Number" como estilos de caractere;
    # preserve-os sob os nomes canônicos que o Page Bible exige.
    for char_name, size in (("Book Title", typo["title_page_title_pt"]),
                            ("Page Number", typo["folio_pt"])):
        st = (doc.styles[char_name] if char_name in doc.styles
              else doc.styles.add_style(char_name, WD_STYLE_TYPE.CHARACTER))
        st.font.name = family
        rpr = st._element.get_or_add_rPr()
        rpr.get_or_add_rFonts().set(qn("w:ascii"), family)
        rpr.rFonts.set(qn("w:hAnsi"), family)
        st.font.size = Pt(size)

    E = lambda *a, **k: ensure_style(doc, *a, family, **k)  # noqa: E731
    E("Front Matter Title", size=typo["front_matter_title_pt"], bold=True,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    E("Copyright", size=typo["copyright_pt"], alignment=WD_ALIGN_PARAGRAPH.LEFT,
      line=typo["copyright_leading_pt"])
    E("Fiction Notice", size=typo["fiction_notice_pt"],
      alignment=WD_ALIGN_PARAGRAPH.LEFT, line=typo["fiction_notice_leading_pt"])
    E("TOC Heading", size=typo["toc_heading_pt"], bold=True,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    toc = E("TOC Entry", size=typo["toc_entry_pt"], alignment=WD_ALIGN_PARAGRAPH.LEFT, line=13)
    toc.paragraph_format.tab_stops.add_tab_stop(
        Inches(typo["toc_tab_position_in"]), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    E("Chapter Number", size=typo["chapter_number_pt"],
      alignment=WD_ALIGN_PARAGRAPH.CENTER, after=10, keep_next=True)
    E("Body First", size=typo["body_pt"], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
      line=typo["body_leading_pt"], first=0)
    E("Body", size=typo["body_pt"], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
      line=typo["body_leading_pt"], first=typo["body_first_line_indent_in"])
    E("Scene Break", size=typo["chapter_number_pt"], alignment=WD_ALIGN_PARAGRAPH.CENTER,
      before=8, after=8, keep_next=True)
    E("Document Header", size=typo["document_block_pt"], bold=True,
      alignment=WD_ALIGN_PARAGRAPH.LEFT, line=typo["document_block_leading_pt"],
      left=typo["document_block_indent_in"], right=typo["document_block_indent_in"])
    E("Document Body", size=typo["document_block_pt"], alignment=WD_ALIGN_PARAGRAPH.LEFT,
      line=typo["document_block_leading_pt"], left=typo["document_block_indent_in"],
      right=typo["document_block_indent_in"])
    E("Image Page", size=1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    E("Header Left", size=typo["running_header_pt"], alignment=WD_ALIGN_PARAGRAPH.LEFT)
    E("Header Right", size=typo["running_header_pt"], alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # Títulos de capítulo nascem em Heading 1 com nível de estrutura 0: mesma
    # aparência do estilo customizado original, mas o Word consegue gerar
    # sumário automático a partir dele sem pós-processamento.
    E("Heading 1", size=typo["chapter_title_pt"], bold=True,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, after=24, keep_next=True, outline=0)

    for built in ("Normal", "Heading 2", "Heading 3"):
        if built in doc.styles:
            st = doc.styles[built]
            st.font.name = family
            rpr = st._element.get_or_add_rPr()
            rpr.get_or_add_rFonts().set(qn("w:ascii"), family)
            rpr.rFonts.set(qn("w:hAnsi"), family)


def configure_page(section, page):
    section.page_width = Inches(page["width_in"])
    section.page_height = Inches(page["height_in"])
    section.top_margin = Inches(page["margin_top_in"])
    section.bottom_margin = Inches(page["margin_bottom_in"])
    # Com mirrorMargins ligado, "left" é a margem interna (lombada).
    section.left_margin = Inches(page["margin_inner_in"])
    section.right_margin = Inches(page["margin_outer_in"])
    section.header_distance = Inches(page["header_distance_in"])
    section.footer_distance = Inches(page["footer_distance_in"])


def clear_part(part):
    for p in part.paragraphs:
        p._element.getparent().remove(p._element)
    part.add_paragraph()


def add_field(paragraph, code, family, display="1", style_id=None, size_pt=8.5):
    """Insere um campo Word (PAGE, PAGEREF). O valor renderizado fica gravado
    para que o arquivo mostre números corretos mesmo antes de o Word recalcular."""
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), code)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    if style_id:
        rstyle = OxmlElement("w:rStyle")
        rstyle.set(qn("w:val"), style_id)
        rpr.append(rstyle)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), family)
    rfonts.set(qn("w:hAnsi"), family)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(round(size_pt * 2))))  # OOXML usa meios-pontos
    rpr.extend([rfonts, sz])
    t = OxmlElement("w:t")
    t.text = display
    r.extend([rpr, t])
    fld.append(r)
    run._element.addnext(fld)
    run._element.getparent().remove(run._element)


def set_page_restart(section, start=1):
    sect = section._sectPr
    pg = sect.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect.append(pg)
    pg.set(qn("w:start"), str(start))
    pg.set(qn("w:fmt"), "decimal")


def clear_page_restart(section):
    pg = section._sectPr.find(qn("w:pgNumType"))
    if pg is not None:
        section._sectPr.remove(pg)


ALL_PARTS = ("header", "even_page_header", "first_page_header",
             "footer", "even_page_footer", "first_page_footer")


def set_body_furniture(section, cfg, book, first_page=False):
    """Cabeçalhos correntes e fólios do corpo. Verso e recto recebem textos
    diferentes; a primeira página do capítulo fica limpa."""
    typo, page = cfg["typography"], cfg["page"]
    family = typo["family"]
    configure_page(section, page)
    clear_page_restart(section)
    for name in ALL_PARTS:
        getattr(section, name).is_linked_to_previous = False
    section.different_first_page_header_footer = first_page
    for name in ALL_PARTS:
        clear_part(getattr(section, name))

    heads = cfg["running_headers"]
    verso = heads["verso"].format(**book)
    recto = heads["recto"].format(**book)
    if heads.get("uppercase", True):
        verso, recto = verso.upper(), recto.upper()

    p_recto = section.header.paragraphs[0]
    p_recto.style = "Header Right"
    set_run_font(p_recto.add_run(recto), family, book["language"], typo["running_header_pt"])

    p_verso = section.even_page_header.paragraphs[0]
    p_verso.style = "Header Left"
    set_run_font(p_verso.add_run(verso), family, book["language"], typo["running_header_pt"])

    f_recto = section.footer.paragraphs[0]
    f_recto.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_field(f_recto, "PAGE", family, style_id="PageNumber", size_pt=typo["folio_pt"])
    f_verso = section.even_page_footer.paragraphs[0]
    f_verso.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_field(f_verso, "PAGE", family, style_id="PageNumber", size_pt=typo["folio_pt"])


def set_blank_furniture(section, cfg):
    """Páginas de imagem, brancas de paridade e matéria frontal: sem cabeçalho
    nem fólio, conforme o Page Bible."""
    configure_page(section, cfg["page"])
    clear_page_restart(section)
    section.different_first_page_header_footer = True
    for name in ALL_PARTS:
        part = getattr(section, name)
        part.is_linked_to_previous = False
        clear_part(part)


def bookmark(paragraph, name, bid):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_rich_text(paragraph, text, family, language, size):
    """Itálico Markdown é a única marcação inline do manuscrito para o leitor."""
    pos = 0
    for m in re.finditer(r"(?<!\*)\*([^*]+)\*(?!\*)", text):
        if m.start() > pos:
            set_run_font(paragraph.add_run(text[pos:m.start()]), family, language, size)
        set_run_font(paragraph.add_run(m.group(1)), family, language, size, italic=True)
        pos = m.end()
    if pos < len(text):
        set_run_font(paragraph.add_run(text[pos:]), family, language, size)


# --------------------------------------------------------------------------
# Construção do documento
# --------------------------------------------------------------------------

def add_image_page(doc, runtime, cfg, chapter, entry, start_type=WD_SECTION.NEW_PAGE):
    img = cfg["images"]
    path = runtime / img["source_dir"] / img["filename_pattern"].format(chapter=chapter)
    if not path.is_file():
        raise SystemExit(f"BUILD FAILED: imagem ausente para o capítulo {chapter}: {path}")

    sec = doc.add_section(start_type)
    set_blank_furniture(sec, cfg)
    p = doc.add_paragraph(style="Image Page")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(img["space_before_pt"])
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    inline = run.add_picture(
        str(path), width=Inches(img["width_in"]), height=Inches(img["height_in"])
    )._inline

    alt = entry.get("alt", "")
    descr = img["alt_text_prefix"].format(chapter=chapter) + alt
    if entry.get("long"):
        descr += img["long_description_prefix"] + entry["long"]
    inline.docPr.set("descr", descr)
    inline.docPr.set("title", f"Ilustração do capítulo {chapter:02d}")
    return sec


def parse_manuscript(path: Path, cfg) -> dict[int, list[str]]:
    if not path.is_file():
        raise SystemExit(f"BUILD FAILED: manuscrito congelado ausente: {path}")
    pattern = re.compile(cfg["manuscript"]["chapter_heading_pattern"])
    ignore = set(cfg["manuscript"].get("ignore_lines", []))
    chapters: dict[int, list[str]] = {}
    current, buf = None, []
    for line in path.read_text(encoding="utf-8").splitlines():
        m = pattern.match(line)
        if m:
            if current is not None:
                chapters[current] = buf
            current, buf = int(m.group(1)), []
        elif current is not None and line.strip() not in ignore:
            buf.append(line)
    if current is not None:
        chapters[current] = buf
    return chapters


def add_front_matter(doc, cfg, book):
    typo, fm = cfg["typography"], cfg["front_matter"]
    family, lang = typo["family"], book["language"]
    title_uc = book["title"].upper()

    sec = doc.sections[0]
    set_blank_furniture(sec, cfg)

    # 1. falsa folha de rosto
    p = doc.add_paragraph(style="Front Matter Title")
    p.paragraph_format.space_before = Pt(fm["half_title_space_before_pt"])
    set_run_font(p.add_run(title_uc), family, lang, typo["half_title_pt"], bold=True)
    doc.add_page_break()

    # 2. verso em branco
    doc.add_paragraph()
    doc.add_page_break()

    # 3. folha de rosto
    p = doc.add_paragraph(style="Front Matter Title")
    p.paragraph_format.space_before = Pt(fm["title_page_space_before_pt"])
    set_run_font(p.add_run(title_uc), family, lang, typo["title_page_title_pt"], bold=True)
    a = doc.add_paragraph()
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(a.add_run(book["author"].upper()), family, lang, typo["title_page_author_pt"])
    doc.add_page_break()

    # 4. copyright — sem ISBN, editora, cidade ou ano inventados
    p = doc.add_paragraph(style="Copyright")
    p.paragraph_format.space_before = Pt(fm["copyright_space_before_pt"])
    set_run_font(p.add_run(fm["copyright_text"].format(**book)), family, lang, typo["copyright_pt"])
    doc.add_page_break()

    # 5. aviso de ficção
    p = doc.add_paragraph(style="Front Matter Title")
    set_run_font(p.add_run(fm["fiction_notice_heading"]), family, lang,
                 typo["front_matter_title_pt"], bold=True)
    q = doc.add_paragraph(style="Fiction Notice")
    set_run_font(q.add_run(fm["fiction_notice_text"].strip()), family, lang,
                 typo["fiction_notice_pt"])
    doc.add_page_break()

    # 6. sumário com campos PAGEREF apontando para os bookmarks dos capítulos
    p = doc.add_paragraph(style="TOC Heading")
    set_run_font(p.add_run(fm["toc_heading"]), family, lang, typo["toc_heading_pt"], bold=True)
    for number in sorted(book["titles"]):
        p = doc.add_paragraph(style="TOC Entry")
        set_run_font(p.add_run(f"{number}. {book['titles'][number]}"), family, lang,
                     typo["toc_entry_pt"])
        p.add_run("\t")
        add_field(p, f"PAGEREF chapter_{number:02d} \\h", family, "0",
                  size_pt=typo["toc_entry_pt"])


def build(runtime: Path) -> int:
    cfg = load_config(runtime)
    book = load_book(runtime)
    placement = load_placement(runtime, book["chapter_count"])
    typo, family, lang = cfg["typography"], cfg["typography"]["family"], book["language"]

    manuscript_path = runtime / cfg["manuscript"]["source"]
    chapters = parse_manuscript(manuscript_path, cfg)

    expected = book["chapter_count"]
    if len(chapters) != expected:
        raise SystemExit(
            f"BUILD FAILED: manuscrito tem {len(chapters)} capítulos, "
            f"BOOK_SPEC declara {expected}."
        )
    missing = [c for c in range(1, expected + 1) if c not in chapters]
    if missing:
        raise SystemExit(f"BUILD FAILED: capítulos ausentes no manuscrito: {missing}")

    doc = Document()
    setup_styles(doc, typo)
    doc.settings.odd_and_even_pages_header_footer = True
    settings = doc.settings._element
    flags = []
    out_cfg = cfg["output"]
    if out_cfg.get("mirror_margins", True):
        flags.append("w:mirrorMargins")
    if out_cfg.get("auto_hyphenation", True):
        flags.append("w:autoHyphenation")
    if out_cfg.get("prevent_image_compression", True):
        flags.append("w:doNotCompressPictures")
    if out_cfg.get("embed_fonts", True):
        flags += ["w:embedTrueTypeFonts", "w:saveSubset"]
    for tag in flags:
        if settings.find(qn(tag)) is None:
            settings.append(OxmlElement(tag))

    doc.core_properties.title = book["title"]
    doc.core_properties.author = book["author"]
    add_front_matter(doc, cfg, book)

    sb = cfg["scene_break"]
    markers = set(sb.get("source_markers", ["* * *"]))
    opening = cfg["chapter_opening"]
    inserted: set[int] = set()
    scene_breaks = 0
    first_chapter = True

    for chapter in range(1, expected + 1):
        entry = placement.get(chapter, {"placement": "OPEN"})
        mode = str(entry.get("placement", "OPEN")).upper()
        anchor = entry.get("anchor")

        # OPEN: página de imagem no verso imediatamente anterior à abertura.
        if mode in PLACEMENT_BEFORE:
            add_image_page(doc, runtime, cfg, chapter, entry, WD_SECTION.EVEN_PAGE)
            inserted.add(chapter)

        start = WD_SECTION.ODD_PAGE if opening.get("start_on_recto", True) else WD_SECTION.NEW_PAGE
        sec = doc.add_section(start)
        set_body_furniture(sec, cfg, book, first_page=True)
        if first_chapter:
            set_page_restart(sec, 1)
            first_chapter = False

        pn = doc.add_paragraph(style="Chapter Number")
        pn.paragraph_format.space_before = Pt(opening["space_before_pt"])
        set_run_font(pn.add_run(opening["number_label"].format(chapter=chapter)),
                     family, lang, typo["chapter_number_pt"])
        pt = doc.add_paragraph(style="Heading 1")
        set_run_font(pt.add_run(book["titles"].get(chapter, f"Capítulo {chapter}")),
                     family, lang, typo["chapter_title_pt"], bold=True)
        bookmark(pt, f"chapter_{chapter:02d}", chapter)

        first_after_break = True
        buf: list[str] = []
        for line in chapters[chapter] + [""]:
            stripped = line.strip()
            if stripped and stripped not in markers:
                buf.append(stripped)
                continue
            if buf:
                text = " ".join(buf)
                p = doc.add_paragraph(style="Body First" if first_after_break else "Body")
                add_rich_text(p, text, family, lang, typo["body_pt"])
                first_after_break = False
                buf = []
                # HINGE/AFTER: a imagem entra somente depois do parágrafo que
                # executa o gatilho declarado; o texto retoma na página seguinte.
                if (mode in PLACEMENT_AFTER and chapter not in inserted
                        and anchor and anchor in text):
                    add_image_page(doc, runtime, cfg, chapter, entry, WD_SECTION.NEW_PAGE)
                    inserted.add(chapter)
                    cont = doc.add_section(WD_SECTION.NEW_PAGE)
                    set_body_furniture(cont, cfg, book, first_page=False)
                    first_after_break = True
            if stripped in markers:
                p = doc.add_paragraph(style="Scene Break")
                set_run_font(p.add_run(sb["glyph"]), family, lang, sb["size_pt"])
                scene_breaks += 1
                first_after_break = True

    never = sorted(set(range(1, expected + 1)) - inserted)
    if never:
        raise SystemExit(
            "BUILD FAILED: imagem não inserida nos capítulos "
            f"{never}. Para HINGE/AFTER, confira se o texto de `anchor` em "
            "layout/IMAGE_PLACEMENT.yaml existe literalmente no manuscrito."
        )

    # Página final par: o Word pode inserir uma branca para satisfazer a paridade.
    tail = doc.add_section(WD_SECTION.EVEN_PAGE)
    set_blank_furniture(tail, cfg)

    out = runtime / cfg["output"]["path"]
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)

    print(f"DOCX OK | {out}")
    print(f"- capítulos: {len(chapters)}")
    print(f"- imagens inseridas: {len(inserted)}")
    print(f"- quebras de cena: {scene_breaks}")
    print(f"- seções: {len(doc.sections)}")
    print(f"- bytes: {out.stat().st_size}")
    print("Auditoria estrutural e render-QA continuam obrigatórios (GATE_KDP).")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Constrói o DOCX de interior KDP a partir do manuscrito congelado."
    )
    parser.add_argument(
        "--runtime", type=Path, default=Path.cwd(),
        help="Raiz do runtime (padrão: diretório atual).",
    )
    args = parser.parse_args()
    runtime = args.runtime.resolve()
    if not (runtime / "book" / "BOOK_SPEC.yaml").is_file():
        print(f"Erro: {runtime} não parece um runtime (falta book/BOOK_SPEC.yaml).",
              file=sys.stderr)
        return 2
    return build(runtime)


if __name__ == "__main__":
    raise SystemExit(main())
